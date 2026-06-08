import io
import re
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def _clean(text: str) -> str:
    """Strip markdown bold/italic markers from text."""
    text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}(.*?)_{1,3}", r"\1", text)
    return text.strip()


def extract_text_from_pdf(file) -> str:
    data = file.read() if hasattr(file, "read") else open(file, "rb").read()
    pdf = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page in pdf:
        pages.append(page.get_text())
    pdf.close()
    return "\n".join(pages)


def extract_resume_text(file) -> str:
    import zipfile
    import io as _io
    from lxml import etree

    data = file.read() if hasattr(file, "read") else open(file, "rb").read()

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    P_TAG = f"{{{W}}}p"
    T_TAG = f"{{{W}}}t"

    try:
        with zipfile.ZipFile(_io.BytesIO(data)) as z:
            with z.open("word/document.xml") as f:
                raw_xml = f.read()
        tree = etree.fromstring(raw_xml)
        lines = []
        for p_el in tree.iter(P_TAG):
            text = "".join(t.text for t in p_el.iter(T_TAG) if t.text).strip()
            if text:
                lines.append(text)
        return "\n".join(lines)
    except Exception:
        # fallback: python-docx
        doc = Document(_io.BytesIO(data))
        P = qn("w:p")
        T = qn("w:t")
        lines = []
        for p_el in doc.element.body.iter(P):
            text = "".join(t.text for t in p_el.iter(T) if t.text).strip()
            if text:
                lines.append(text)
        return "\n".join(lines)


def extract_resume_sections(file) -> dict:
    """Extract text grouped by section headings."""
    doc = Document(file)
    sections = {}
    current_section = "header"
    current_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower()
        is_heading = "heading" in style_name or para.runs and any(
            r.bold for r in para.runs if r.text.strip()
        ) and len(text) < 60

        if is_heading and current_lines:
            sections[current_section] = "\n".join(current_lines)
            current_section = text
            current_lines = []
        else:
            current_lines.append(text)

    if current_lines:
        sections[current_section] = "\n".join(current_lines)

    return sections


def build_output_docx(tailored: dict, specialist_name: str, template_path: str = None) -> bytes:
    """
    Build a polished .docx from tailored sections.
    If template_path is provided and has <<PLACEHOLDERS>>, fills them in.
    Otherwise generates a clean document from scratch.
    """
    if template_path:
        try:
            return _fill_template(template_path, tailored, specialist_name)
        except Exception:
            pass

    return _build_from_scratch(tailored, specialist_name)


def _fill_template(template_path: str, tailored: dict, specialist_name: str) -> bytes:
    doc = Document(template_path)

    placeholder_map = {
        "<<NAME>>": tailored.get("name", specialist_name),
        "<<SUMMARY>>": tailored.get("summary", ""),
        "<<SKILLS>>": tailored.get("skills", ""),
        "<<EXPERIENCE>>": tailored.get("experience", ""),
        "<<EDUCATION>>": tailored.get("education", ""),
        "<<CONTACTS>>": tailored.get("contacts", ""),
    }

    for para in doc.paragraphs:
        for placeholder, value in placeholder_map.items():
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in placeholder_map.items():
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_section_heading(doc: Document, text: str):
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 2"]
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Add bottom border to heading
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_text(doc: Document, text: str, size: int = 10, bold: bool = False,
              color: RGBColor = None, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0C0C0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_from_scratch(tailored: dict, specialist_name: str) -> bytes:
    # Colors matching corporate template
    NAVY = RGBColor(0x1F, 0x38, 0x64)
    GREEN = RGBColor(0x00, 0xA8, 0x58)
    GRAY = RGBColor(0x4D, 0x4D, 0x50)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Name
    p = doc.add_paragraph()
    r = p.add_run(_clean(tailored.get("name", specialist_name)))
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY

    # Role + City
    role = tailored.get("role", "")
    city = tailored.get("city", "")
    if role:
        p = doc.add_paragraph()
        r = p.add_run(_clean(role))
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = NAVY
    if city:
        p = doc.add_paragraph()
        r = p.add_run(_clean(city))
        r.font.size = Pt(14)
        r.font.color.rgb = NAVY

    doc.add_paragraph()

    # Summary
    summary = tailored.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        r = p.add_run(_clean(summary))
        r.font.size = Pt(10)
        r.font.color.rgb = GRAY

    doc.add_paragraph()

    # Skills block
    skills = tailored.get("skills", "")
    if skills:
        p = doc.add_paragraph()
        r_label = p.add_run("Ключевые навыки: ")
        r_label.bold = True
        r_label.font.size = Pt(10)
        r_label.font.color.rgb = GRAY
        r_val = p.add_run(_clean(skills.replace("\n", ", ").strip(", ")))
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = GRAY
        doc.add_paragraph()

    # Work experience — structured per project
    projects = tailored.get("projects", [])
    for idx, proj in enumerate(projects):
        # Role + Company (green header)
        p = doc.add_paragraph()
        r = p.add_run(_clean(proj.get("role_company", "")))
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = GREEN

        # Dates + type
        dates = proj.get("dates", "")
        proj_type = proj.get("type", "")
        if dates or proj_type:
            p = doc.add_paragraph()
            r = p.add_run(f"({_clean(dates)})" if not proj_type else f"({_clean(dates)})  ({_clean(proj_type)})")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY

        # Project description
        description = proj.get("description", "")
        if description:
            p = doc.add_paragraph()
            r_label = p.add_run("Описание: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
            r_label.font.color.rgb = GRAY
            r_val = p.add_run(_clean(description))
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = GRAY

        doc.add_paragraph()

        # Tasks
        tasks = proj.get("tasks", [])
        if tasks:
            p = doc.add_paragraph()
            r = p.add_run("Задачи:")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY
            doc.add_paragraph()
            for task in tasks:
                bp = doc.add_paragraph(style="List Bullet")
                br = bp.add_run(_clean(task.lstrip("•-– ")))
                br.font.size = Pt(10)
                br.font.color.rgb = GRAY

        doc.add_paragraph()

        # Achievements
        achievements = proj.get("achievements", [])
        if achievements:
            p = doc.add_paragraph()
            r = p.add_run("Достижения:")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = GRAY
            doc.add_paragraph()
            for ach in achievements:
                bp = doc.add_paragraph(style="List Bullet")
                br = bp.add_run(_clean(ach.lstrip("•-– ")))
                br.font.size = Pt(10)
                br.font.color.rgb = GRAY

        doc.add_paragraph()

        # Tech stack
        stack = proj.get("stack", "")
        if stack:
            p = doc.add_paragraph()
            r_label = p.add_run("Используемый стек: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
            r_label.font.color.rgb = GRAY
            r_val = p.add_run(_clean(stack))
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = GRAY

        # Team
        team = proj.get("team", "")
        if team:
            p = doc.add_paragraph()
            r_label = p.add_run("Команда: ")
            r_label.bold = True
            r_label.font.size = Pt(10)
            r_label.font.color.rgb = GRAY
            r_val = p.add_run(_clean(team))
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = GRAY

        # Separator between projects
        if idx < len(projects) - 1:
            doc.add_paragraph()
            _add_separator(doc)
            doc.add_paragraph()

    # Education
    education = tailored.get("education", "")
    if education:
        doc.add_paragraph()
        _add_separator(doc)
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Образование")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = GREEN
        for line in education.split("\n"):
            line = _clean(line.strip())
            if line:
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(10)
                r.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
